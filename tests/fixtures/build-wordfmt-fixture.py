#!/usr/bin/env python3
"""Build a deterministic DOCX fixture for WordFmt end-to-end tests."""

from pathlib import Path
from io import BytesIO
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches
from PIL import Image, ImageDraw


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    run.append(run_properties)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build(output_path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    doc.add_heading("CHƯƠNG 1. GIỚI THIỆU TỔNG QUAN", level=1)
    doc.add_heading("1.1 Mục tiêu của đề tài", level=2)

    body = doc.add_paragraph()
    body.add_run('Đoạn văn có "ngoặc kép thẳng", khoảng 2–4 và một mệnh đề — cần đổi em dash. ')
    body.add_run("Cụm in đậm cần được giữ.").bold = True
    body.add_run(" Cụm in nghiêng cần được giữ.").italic = True

    doc.add_paragraph("Danh sách dấu tròn", style="List Bullet")
    doc.add_paragraph("Danh sách đánh số", style="List Number")
    doc.add_paragraph("- Tiểu mục dùng dấu trừ")
    doc.add_paragraph("+ Tiểu mục dùng dấu cộng")

    doc.add_heading("1.1.1 Thiết kế tổng thể", level=3)
    doc.add_heading("1.1.1.1 Thành phần", level=4)

    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "STT"
    table.rows[0].cells[1].text = "Nội dung"
    table.rows[0].cells[2].text = "Kết quả"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "Kiểm tra bảng"
    table.rows[1].cells[2].text = "Đạt"
    table_caption = doc.add_paragraph("Bảng 1-1: Bảng kiểm tra định dạng")
    table_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    figure_caption = doc.add_paragraph("Hình 1-1: Hình kiểm tra định dạng")
    figure_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image = Image.new("RGB", (640, 320), "white")
    canvas = ImageDraw.Draw(image)
    canvas.rectangle((20, 20, 620, 300), outline="black", width=4)
    canvas.text((170, 145), "WORDFMT FIGURE", fill="black")
    image_buffer = BytesIO()
    image.save(image_buffer, format="PNG")
    image_buffer.seek(0)
    figure = doc.add_paragraph()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.add_run().add_picture(image_buffer, width=Inches(4.5))

    doc.add_page_break()
    doc.add_heading("CHƯƠNG 2. KẾT QUẢ", level=1)
    doc.add_heading("2.1 Kết quả kiểm tra", level=2)
    doc.add_paragraph("Nội dung chương hai dùng để kiểm tra header theo section.")

    doc.add_page_break()
    doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    reference = doc.add_paragraph()
    reference.add_run("[1] Tác giả. (2026). ")
    reference.add_run("Tên sách cần in nghiêng").italic = True
    reference.add_run(". Truy cập tại ")
    add_hyperlink(reference, "https://example.com/reference", "https://example.com/reference")
    reference.add_run(".")

    second_reference = doc.add_paragraph()
    second_reference.add_run("[2] Tài liệu có liên kết thứ hai: ")
    add_hyperlink(second_reference, "Trang tài liệu", "https://example.org/docs")

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("usage: build-wordfmt-fixture.py OUTPUT.docx")
    build(sys.argv[1])
